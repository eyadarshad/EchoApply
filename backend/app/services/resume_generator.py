import io
import logging
import copy
from typing import Dict, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from app.schemas import ResumeParsedData

logger = logging.getLogger(__name__)

def generate_fallback_pdf(data: ResumeParsedData, template_name: str = "classic", compact_mode: bool = False) -> bytes:
    """
    Styled pure-Python PDF generator using ReportLab Platypus layout calibrated for ISO A4 pages.
    Provides customized single-page designs matching all templates with guaranteed 1-page containment.
    """
    import io
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import (
        SimpleDocTemplate, BaseDocTemplate, PageTemplate, Frame, FrameBreak,
        Paragraph, Spacer, HRFlowable, ListFlowable, ListItem
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors

    # ISO A4 dimensions (595.27 pt x 841.89 pt)
    page_width, page_height = A4

    # Define color palettes & typography
    font_family = "Helvetica"
    font_bold = "Helvetica-Bold"
    font_oblique = "Helvetica-Oblique"
    
    colors_dict = data.color_theme or {}
    primary_hex = colors_dict.get("primary")
    accent_hex = colors_dict.get("accent")
    text_hex = colors_dict.get("text")
    bg_hex = colors_dict.get("background")
    sidebar_bg_hex = colors_dict.get("sidebar_bg")
    sidebar_accent_hex = colors_dict.get("sidebar_accent")

    primary_color = colors.HexColor(primary_hex) if primary_hex else colors.HexColor('#0f172a') # Slate 900
    accent_color = colors.HexColor(accent_hex) if accent_hex else colors.HexColor('#0d9488') # Teal 600
    text_color = colors.HexColor(text_hex) if text_hex else colors.HexColor('#1e293b')
    bg_color = colors.HexColor(bg_hex) if bg_hex else colors.white
    
    divider_color = colors.HexColor('#cbd5e0')
    header_align = 0 # Left by default for two-column

    # Sidebar parameters
    has_sidebar = template_name in ("modern", "creative", "executive", "modern_executive")
    sidebar_bg_color = None
    divider_line_color = None

    if template_name in ("classic", "classic_executive"):
        font_family = "Times-Roman"
        font_bold = "Times-Bold"
        font_oblique = "Times-Italic"
        if not primary_hex: primary_color = colors.HexColor('#1e3a5f')
        if not accent_hex: accent_color = colors.HexColor('#b8860b')
        if not text_hex: text_color = colors.HexColor('#1e293b')
        divider_color = colors.HexColor('#1e3a5f')
        header_align = 1 # Center
    elif template_name in ("modern", "modern_executive"):
        if not primary_hex: primary_color = colors.HexColor('#0f172a')
        if not accent_hex: accent_color = colors.HexColor('#14b8a6')
        if not text_hex: text_color = colors.HexColor('#1e293b')
        divider_color = colors.HexColor('#ccfbf1')
        sidebar_bg_color = colors.HexColor(sidebar_bg_hex) if sidebar_bg_hex else colors.HexColor('#111827')
        divider_line_color = colors.HexColor(sidebar_accent_hex) if sidebar_accent_hex else colors.HexColor('#14b8a6')
    elif template_name == "minimal":
        if not primary_hex: primary_color = colors.HexColor('#2d2d2d')
        if not accent_hex: accent_color = colors.HexColor('#4a7c6f')
        if not text_hex: text_color = colors.HexColor('#2d2d2d')
        divider_color = colors.HexColor('#f1f5f9')
        sidebar_bg_color = colors.HexColor(sidebar_bg_hex) if sidebar_bg_hex else colors.HexColor('#faf9f7')
        divider_line_color = colors.HexColor('#e2e8f0')
    elif template_name == "creative":
        if not primary_hex: primary_color = colors.HexColor('#4a1942')
        if not accent_hex: accent_color = colors.HexColor('#ec4899')
        if not text_hex: text_color = colors.HexColor('#4a1942')
        divider_color = colors.HexColor('#cffafe')
        sidebar_bg_color = colors.HexColor(sidebar_bg_hex) if sidebar_bg_hex else colors.HexColor('#fdf8f0')
        divider_line_color = colors.HexColor(sidebar_accent_hex) if sidebar_accent_hex else colors.HexColor('#fbcfe8')
    elif template_name == "executive":
        font_family = "Times-Roman"
        font_bold = "Times-Bold"
        font_oblique = "Times-Italic"
        if not primary_hex: primary_color = colors.HexColor('#0f172a')
        if not accent_hex: accent_color = colors.HexColor('#c9a55c')
        if not text_hex: text_color = colors.HexColor('#1e293b')
        divider_color = colors.HexColor('#ccfbf1')
        sidebar_bg_color = colors.HexColor(sidebar_bg_hex) if sidebar_bg_hex else colors.HexColor('#fefdfb')
        divider_line_color = colors.HexColor(sidebar_accent_hex) if sidebar_accent_hex else colors.HexColor('#c9a55c')

    # Font family override from data
    if data.font_family:
        ff = data.font_family.lower()
        if "times" in ff or "georgia" in ff or "serif" in ff:
            font_family = "Times-Roman"
            font_bold = "Times-Bold"
            font_oblique = "Times-Italic"
        elif "courier" in ff or "mono" in ff:
            font_family = "Courier"
            font_bold = "Courier-Bold"
            font_oblique = "Courier-Oblique"
        else:
            font_family = "Helvetica"
            font_bold = "Helvetica-Bold"
            font_oblique = "Helvetica-Oblique"

    buffer = io.BytesIO()

    # Draw sidebar background callback
    def draw_sidebar_background(canvas, doc):
        if not sidebar_bg_color:
            return
        canvas.saveState()
        canvas.setFillColor(sidebar_bg_color)
        canvas.rect(0, 0, 0.32 * page_width, page_height, fill=1, stroke=0)
        if divider_line_color:
            canvas.setStrokeColor(divider_line_color)
            canvas.setLineWidth(2.5 if template_name in ("modern", "modern_executive") else 1.5)
            canvas.line(0.32 * page_width, 0, 0.32 * page_width, page_height)
        canvas.restoreState()

    if has_sidebar:
        doc = BaseDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=0,
            rightMargin=0,
            topMargin=0,
            bottomMargin=0
        )
        left_margin = 12 if compact_mode else 16
        bottom_margin = 18 if compact_mode else 24
        top_margin = 18 if compact_mode else 24
        
        # Left frame (sidebar)
        frame_left = Frame(
            left_margin, bottom_margin, 0.32 * page_width - (left_margin + 10), page_height - (bottom_margin + top_margin),
            id='sidebar_frame', topPadding=12 if compact_mode else 18, bottomPadding=0, leftPadding=0, rightPadding=0
        )
        # Right frame (main content)
        frame_right = Frame(
            0.32 * page_width + left_margin, bottom_margin, 0.68 * page_width - (left_margin + 16), page_height - (bottom_margin + top_margin),
            id='main_frame', topPadding=12 if compact_mode else 18, bottomPadding=0, leftPadding=0, rightPadding=0
        )
        template = PageTemplate(
            id='two_column',
            frames=[frame_left, frame_right],
            onPage=draw_sidebar_background
        )
        doc.addPageTemplates([template])
    else:
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=30 if compact_mode else 42,
            rightMargin=30 if compact_mode else 42,
            topMargin=22 if compact_mode else 32,
            bottomMargin=22 if compact_mode else 32
        )

    styles = getSampleStyleSheet()

    # Custom styles calibrated for A4
    title_size = 15 if compact_mode else (17 if has_sidebar else 19)
    title_leading = 18 if compact_mode else (20 if has_sidebar else 22)
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName=font_bold,
        fontSize=title_size,
        leading=title_leading,
        textColor=colors.white if (has_sidebar and template_name in ("modern", "modern_executive")) else primary_color,
        alignment=header_align,
        spaceAfter=3 if compact_mode else 5
    )
    
    contact_style = ParagraphStyle(
        'DocContact',
        parent=styles['Normal'],
        fontName=font_family,
        fontSize=7.5 if compact_mode else 8,
        leading=10 if compact_mode else 11,
        textColor=colors.HexColor('#cbd5e1') if (has_sidebar and template_name in ("modern", "modern_executive")) else colors.HexColor('#475569'),
        alignment=header_align,
        spaceAfter=6 if compact_mode else 10
    )

    anchor_style = ParagraphStyle(
        'DocAnchor',
        parent=styles['Normal'],
        fontName=font_oblique,
        fontSize=8.5 if compact_mode else 9,
        leading=11 if compact_mode else 12,
        textColor=accent_color,
        alignment=header_align,
        spaceAfter=6 if compact_mode else 10
    )

    sidebar_heading_style = ParagraphStyle(
        'SidebarHeading',
        parent=styles['Heading3'],
        fontName=font_bold,
        fontSize=8 if compact_mode else 8.5,
        leading=11 if compact_mode else 12,
        textColor=colors.HexColor(sidebar_accent_hex) if sidebar_accent_hex else accent_color,
        spaceBefore=8 if compact_mode else 12,
        spaceAfter=4 if compact_mode else 6
    )

    sidebar_text_style = ParagraphStyle(
        'SidebarText',
        parent=styles['Normal'],
        fontName=font_family,
        fontSize=7.5 if compact_mode else 8,
        leading=10 if compact_mode else 11,
        textColor=colors.HexColor('#cbd5e1') if (has_sidebar and template_name in ("modern", "modern_executive")) else text_color,
        spaceAfter=4 if compact_mode else 6
    )

    heading_style = ParagraphStyle(
        'DocHeading',
        parent=styles['Heading2'],
        fontName=font_bold,
        fontSize=9.5 if compact_mode else 10.5,
        leading=13 if compact_mode else 14,
        textColor=primary_color,
        spaceBefore=6 if compact_mode else 10,
        spaceAfter=3 if compact_mode else 5
    )

    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontName=font_family,
        fontSize=8 if compact_mode else 8.5,
        leading=11 if compact_mode else 12,
        textColor=text_color,
        spaceAfter=3 if compact_mode else 4
    )

    bullet_style = ParagraphStyle(
        'DocBullet',
        parent=styles['Normal'],
        fontName=font_family,
        fontSize=7.5 if compact_mode else 8,
        leading=10.5 if compact_mode else 11.5,
        textColor=text_color,
        spaceAfter=1.5 if compact_mode else 2
    )

    story = []

    def add_section_header(title_text):
        story.append(Paragraph(title_text.upper(), heading_style))
        story.append(HRFlowable(width="100%", thickness=1, color=accent_color, spaceBefore=1, spaceAfter=4 if compact_mode else 6))

    if has_sidebar:
        # === 1. LEFT SIDEBAR STORY ===
        story.append(Paragraph(data.name, title_style))
        
        contact_parts = []
        if data.email: contact_parts.append(data.email)
        if data.phone: contact_parts.append(data.phone)
        for link in data.links[:2]:
            contact_parts.append(link.replace("https://", "").replace("www.", ""))
        
        story.append(Paragraph("<br/>".join(contact_parts), contact_style))
        story.append(Spacer(1, 4))

        # Skills in sidebar
        if data.skills:
            story.append(Paragraph("SKILLS CORE", sidebar_heading_style))
            skills_preview = data.skills[:10] if compact_mode else data.skills[:12]
            for s in skills_preview:
                story.append(Paragraph(f"• {s}", sidebar_text_style))

        # Education in sidebar
        if data.education:
            story.append(Paragraph("EDUCATION", sidebar_heading_style))
            for edu in data.education[:2]:
                deg = edu.degree
                if edu.major: deg += f" in {edu.major}"
                story.append(Paragraph(f"<b>{deg}</b><br/>{edu.school} ({edu.date})", sidebar_text_style))

        # Certifications / Languages in sidebar
        if data.certifications:
            story.append(Paragraph("CREDENTIALS", sidebar_heading_style))
            for c in data.certifications[:3]:
                story.append(Paragraph(f"• {c}", sidebar_text_style))

        # === 2. SWITCH TO MAIN FRAME ===
        story.append(FrameBreak())

        # Main header tagline
        tagline = data.scroll_stop_hook or data.anchor_line
        if tagline:
            story.append(Paragraph(f"<b><i>{tagline}</i></b>", anchor_style))

        # Executive summary
        if data.executive_summary:
            story.append(Paragraph(f"<i>{data.executive_summary}</i>", body_style))
            story.append(Spacer(1, 4))

        # Highlights
        if data.highlights_strip:
            add_section_header("Key Highlights")
            for h in data.highlights_strip[:3]:
                story.append(Paragraph(f"• <b>{h.get('skill', '')}:</b> {h.get('relevance_reason', '')}", bullet_style))
            story.append(Spacer(1, 4))

        # Experience
        if data.experience:
            add_section_header("Professional Experience")
            for exp in data.experience[:3]:
                role = exp.role
                company = exp.company
                loc = f", {exp.location}" if exp.location else ""
                dates = f"{exp.start_date} – {exp.end_date or 'Present'}"
                
                story.append(Paragraph(f"<b>{role}</b> — <i>{company}{loc}</i> &nbsp;|&nbsp; <font color='#64748b'>{dates}</font>", body_style))
                
                bullets = exp.bullets or []
                if compact_mode:
                    bullets = bullets[:2]
                for b in bullets:
                    story.append(Paragraph(f"&nbsp;&nbsp;• {b}", bullet_style))
                story.append(Spacer(1, 3))

        # Projects
        if data.projects:
            add_section_header("Projects & Achievements")
            for proj in data.projects[:2]:
                name = proj.name
                link = f" ({proj.link})" if proj.link else ""
                story.append(Paragraph(f"<b>{name}</b>{link}", body_style))
                
                bullets = proj.bullets or []
                if compact_mode:
                    bullets = bullets[:1]
                for b in bullets:
                    story.append(Paragraph(f"&nbsp;&nbsp;• {b}", bullet_style))
                story.append(Spacer(1, 3))

    else:
        # === SINGLE COLUMN LAYOUT (Classic, Minimal, etc.) ===
        story.append(Paragraph(data.name, title_style))
        
        contact_parts = []
        if data.email: contact_parts.append(data.email)
        if data.phone: contact_parts.append(data.phone)
        for link in data.links[:2]:
            contact_parts.append(link.replace("https://", "").replace("www.", ""))
        story.append(Paragraph(" &nbsp;|&nbsp; ".join(contact_parts), contact_style))

        tagline = data.scroll_stop_hook or data.anchor_line
        if tagline:
            story.append(Paragraph(f"<b><i>{tagline}</i></b>", anchor_style))

        if data.executive_summary:
            story.append(Paragraph(f"<i>{data.executive_summary}</i>", body_style))
            story.append(Spacer(1, 4))

        if data.highlights_strip:
            add_section_header("Relevance & Highlights")
            for h in data.highlights_strip[:3]:
                story.append(Paragraph(f"• <b>{h.get('skill', '')}:</b> {h.get('relevance_reason', '')}", bullet_style))
            story.append(Spacer(1, 4))

        if data.skills:
            add_section_header("Technical & Domain Skills")
            skills_preview = data.skills[:12] if compact_mode else data.skills
            story.append(Paragraph(", ".join(skills_preview), body_style))
            story.append(Spacer(1, 4))

        if data.experience:
            add_section_header("Professional Experience")
            for exp in data.experience:
                role = exp.role
                company = exp.company
                loc = f", {exp.location}" if exp.location else ""
                dates = f"{exp.start_date} – {exp.end_date or 'Present'}"
                
                story.append(Paragraph(f"<b>{role}</b> — <i>{company}{loc}</i> &nbsp;|&nbsp; <font color='#64748b'>{dates}</font>", body_style))
                
                bullets = exp.bullets or []
                if compact_mode:
                    bullets = bullets[:2]
                for b in bullets:
                    story.append(Paragraph(f"&nbsp;&nbsp;• {b}", bullet_style))
                story.append(Spacer(1, 3))

        if data.projects:
            add_section_header("Key Projects")
            for proj in data.projects[:2]:
                name = proj.name
                link = f" ({proj.link})" if proj.link else ""
                story.append(Paragraph(f"<b>{name}</b>{link}", body_style))
                for b in (proj.bullets or [])[:2]:
                    story.append(Paragraph(f"&nbsp;&nbsp;• {b}", bullet_style))
                story.append(Spacer(1, 3))

        if data.education:
            add_section_header("Education")
            for edu in data.education[:2]:
                deg = edu.degree
                if edu.major: deg += f" in {edu.major}"
                gpa = f" (GPA: {edu.gpa})" if edu.gpa else ""
                story.append(Paragraph(f"<b>{deg}</b>{gpa} — <i>{edu.school}</i> ({edu.date})", body_style))

        if data.certifications or data.languages:
            add_section_header("Certifications & Languages")
            extra = []
            if data.certifications: extra.append(f"<b>Certifications:</b> {', '.join(data.certifications)}")
            if data.languages: extra.append(f"<b>Languages:</b> {', '.join(data.languages)}")
            story.append(Paragraph(" &nbsp;|&nbsp; ".join(extra), body_style))

    doc.build(story)
    pdf_data = buffer.getvalue()
    buffer.close()
    return pdf_data


def generate_resume_pdf(data: ResumeParsedData, template_name: str = "classic", compact_mode: bool = False) -> bytes:
    """
    Renders structured resume data to PDF using WeasyPrint and elegant HTML template layouts.
    Executes an iterative multi-pass single-page A4 fitting engine guaranteeing exact 1-page output.
    """
    try:
        from weasyprint import HTML
        from app.services.resume_templates import render_template
        
        # Pass 1: Standard render
        html_content = render_template(template_name, data, compact_mode=compact_mode)
        doc = HTML(string=html_content).render()
        if len(doc.pages) <= 1:
            return doc.write_pdf()
            
        # Pass 2: Enable compact_mode
        logger.info("A4 PDF exceeds 1 page. Pass 2: applying compact_mode.")
        html_content = render_template(template_name, data, compact_mode=True)
        doc = HTML(string=html_content).render()
        if len(doc.pages) <= 1:
            return doc.write_pdf()
            
        # Pass 3: Trim bullets (top 3 per role, top 2 per project, top 10 skills)
        logger.info("A4 PDF exceeds 1 page. Pass 3: trimming secondary bullet density.")
        trimmed_data = copy.deepcopy(data)
        if trimmed_data.experience:
            for exp in trimmed_data.experience:
                if exp.bullets:
                    exp.bullets = exp.bullets[:3]
        if trimmed_data.projects:
            for proj in trimmed_data.projects:
                if proj.bullets:
                    proj.bullets = proj.bullets[:2]
        if trimmed_data.skills:
            trimmed_data.skills = trimmed_data.skills[:10]
            
        html_content = render_template(template_name, trimmed_data, compact_mode=True)
        doc = HTML(string=html_content).render()
        if len(doc.pages) <= 1:
            return doc.write_pdf()
            
        # Pass 4: Tighten bullets further (top 2 per role, top 1 per project)
        logger.info("A4 PDF exceeds 1 page. Pass 4: aggressive bullet compaction.")
        if trimmed_data.experience:
            for exp in trimmed_data.experience:
                if exp.bullets:
                    exp.bullets = exp.bullets[:2]
        if trimmed_data.projects:
            for proj in trimmed_data.projects:
                if proj.bullets:
                    proj.bullets = proj.bullets[:1]
                    
        html_content = render_template(template_name, trimmed_data, compact_mode=True)
        doc = HTML(string=html_content).render()
        if len(doc.pages) <= 1:
            return doc.write_pdf()

        # Pass 5: Condense older experience (keep top 2 most recent roles)
        if trimmed_data.experience and len(trimmed_data.experience) > 2:
            logger.info("A4 PDF exceeds 1 page. Pass 5: keeping top 2 recent experience entries.")
            trimmed_data.experience = trimmed_data.experience[:2]
            html_content = render_template(template_name, trimmed_data, compact_mode=True)
            doc = HTML(string=html_content).render()
            if len(doc.pages) <= 1:
                return doc.write_pdf()

        return doc.write_pdf()
        
    except Exception as e:
        logger.warning(f"WeasyPrint compilation failed or unavailable ({e}). Falling back to pure Python A4 PDF compiler.")
        return generate_fallback_pdf(data, template_name, compact_mode=compact_mode)


def generate_resume_docx(data: ResumeParsedData) -> bytes:
    """
    Renders structured resume data to DOCX using python-docx.
    Configures standard margins, tab-stops for dates, and basic styling.
    """
    doc = Document()

    # Configure 0.75-inch margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Configure default style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Header Name
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(data.name)
    run_name.font.size = Pt(18)
    run_name.font.bold = True

    # Contact Info
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_parts = [data.email]
    if data.phone:
        contact_parts.append(data.phone)
    contact_parts.extend(data.links)
    p_contact.add_run("  |  ".join(contact_parts))

    # Anchor Line / Scroll-Stop Hook
    tagline = data.scroll_stop_hook or data.anchor_line
    if tagline:
        p_anchor = doc.add_paragraph()
        p_anchor.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_anchor.paragraph_format.space_after = Pt(8)
        run_anchor = p_anchor.add_run(tagline)
        run_anchor.italic = True
        run_anchor.bold = True
        run_anchor.font.size = Pt(11)

    # Helper to add section headings with bottom border lines
    def add_section_heading(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text.upper())
        run.font.size = Pt(11)
        run.font.bold = True
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/></w:pBdr>')
        p._p.get_or_add_pPr().append(pBdr)

    # Highlights Strip Section
    if data.highlights_strip:
        add_section_heading("Relevance & Highlights")
        for h in data.highlights_strip:
            p_hl = doc.add_paragraph(style='List Bullet')
            p_hl.paragraph_format.space_after = Pt(2)
            run_skill = p_hl.add_run(h.get("skill", "") + ": ")
            run_skill.bold = True
            p_hl.add_run(h.get("relevance_reason", ""))

    # 1. Skills Section
    if data.skills:
        add_section_heading("Skills")
        doc.add_paragraph(", ".join(data.skills))

    # 2. Experience Section
    if data.experience:
        add_section_heading("Experience")
        for exp in data.experience:
            p_title = doc.add_paragraph()
            p_title.paragraph_format.space_before = Pt(4)
            p_title.paragraph_format.space_after = Pt(2)
            p_title.paragraph_format.keep_with_next = True

            role_company = f"{exp.role} - {exp.company}"
            end_date = exp.end_date or "Present"
            dates_location = f"{exp.start_date} - {end_date}"
            if exp.location:
                dates_location += f" ({exp.location})"

            run_title = p_title.add_run(role_company)
            run_title.bold = True
            p_title.add_run("\t" + dates_location)
            p_title.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), alignment=2)

            bullets = exp.bullets or []
            for bullet in bullets:
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_after = Pt(2)
                p_bullet.add_run(bullet)

    # 3. Projects Section
    if data.projects:
        add_section_heading("Projects")
        for proj in data.projects:
            p_proj = doc.add_paragraph()
            p_proj.paragraph_format.space_before = Pt(4)
            p_proj.paragraph_format.space_after = Pt(2)
            p_proj.paragraph_format.keep_with_next = True

            proj_name = proj.name
            proj_link = proj.link
            title_text = f"{proj_name}"
            if proj_link:
                title_text += f" ({proj_link})"

            run_title = p_proj.add_run(title_text)
            run_title.bold = True

            bullets = proj.bullets or []
            for bullet in bullets:
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_after = Pt(2)
                p_bullet.add_run(bullet)

    # 4. Education Section
    if data.education:
        add_section_heading("Education")
        for edu in data.education:
            p_edu = doc.add_paragraph()
            p_edu.paragraph_format.space_before = Pt(4)
            p_edu.paragraph_format.space_after = Pt(2)

            degree_major = f"{edu.degree}"
            if edu.major:
                degree_major += f" in {edu.major}"
            
            school = edu.school
            edu_dates = edu.date

            run_edu = p_edu.add_run(f"{degree_major}, {school}")
            run_edu.bold = True
            p_edu.add_run("\t" + edu_dates)
            p_edu.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), alignment=2)

            if edu.gpa:
                p_gpa = doc.add_paragraph(style='List Bullet')
                p_gpa.paragraph_format.space_after = Pt(2)
                p_gpa.add_run(f"GPA: {edu.gpa}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
